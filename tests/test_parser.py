import pytest
from pathlib import Path
from writer_agent.analysis.parser import parse_document


def test_parse_txt(tmp_path):
    f = tmp_path / "sample.txt"
    f.write_text("Dark text here.", encoding="utf-8")
    result = parse_document(f)
    assert result == "Dark text here."


def test_parse_md(tmp_path):
    f = tmp_path / "sample.md"
    f.write_text("# Chapter 1\n\nDark **text** here.", encoding="utf-8")
    result = parse_document(f)
    assert "Chapter 1" in result
    assert "text" in result


def test_parse_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Dark paragraph one.")
    doc.add_paragraph("Dark paragraph two.")
    f = tmp_path / "sample.docx"
    doc.save(str(f))
    result = parse_document(f)
    assert "Dark paragraph one." in result
    assert "Dark paragraph two." in result


def test_parse_unsupported(tmp_path):
    f = tmp_path / "sample.unsupported"
    f.write_text("stuff", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported format"):
        parse_document(f)
