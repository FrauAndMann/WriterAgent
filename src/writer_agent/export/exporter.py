from pathlib import Path
from writer_agent.db.repositories import ChapterRepo


class Exporter:
    def __init__(self, chapter_repo: ChapterRepo):
        self.repo = chapter_repo

    def to_markdown(self, project_id: int, output_path: Path, title: str = ""):
        chapters = self.repo.list_by_project(project_id)
        lines = [f"# {title}\n"] if title else []
        for ch in sorted(chapters, key=lambda c: c["chapter_number"]):
            ch_title = ch.get("title") or f"Глава {ch['chapter_number']}"
            lines.append(f"\n## {ch_title}\n")
            lines.append(ch["full_text"])
            lines.append("\n---\n")
        output_path.write_text("\n".join(lines), encoding="utf-8")

    def to_txt(self, project_id: int, output_path: Path, title: str = ""):
        chapters = self.repo.list_by_project(project_id)
        lines = [f"{title}\n{'=' * 40}\n"] if title else []
        for ch in sorted(chapters, key=lambda c: c["chapter_number"]):
            ch_title = ch.get("title") or f"Глава {ch['chapter_number']}"
            lines.append(f"\n{ch_title}\n")
            lines.append(ch["full_text"])
            lines.append("\n" + "-" * 40 + "\n")
        output_path.write_text("\n".join(lines), encoding="utf-8")

    def to_docx(self, project_id: int, output_path: Path, title: str = ""):
        from docx import Document
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        chapters = self.repo.list_by_project(project_id)
        for ch in sorted(chapters, key=lambda c: c["chapter_number"]):
            ch_title = ch.get("title") or f"Глава {ch['chapter_number']}"
            doc.add_heading(ch_title, level=1)
            for paragraph in ch["full_text"].split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        doc.save(str(output_path))
